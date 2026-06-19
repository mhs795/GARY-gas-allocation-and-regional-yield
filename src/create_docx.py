from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_docs():
    doc = Document()

    # Title
    title = doc.add_heading('Australian Nodal Gas Market Optimization Model', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph('Strategic Multi-Year Dispatch and Infrastructure Planning Framework (2025–2050)')

    # Introduction
    doc.add_heading('1. Executive Summary', level=1)
    doc.add_paragraph(
        'This model is a mathematical optimization framework designed to simulate the Australian east coast gas market\'s '
        'transition through to 2050. It utilizes Nodal Least-Cost Dispatch logic to solve for the most efficient gas '
        'production, pipeline transmission, and infrastructure expansion decisions, while accounting for the structural '
        'shifts identified in the AEMO Gas Statement of Opportunities (GSOO).'
    )

    # Background
    doc.add_heading('2. Model Background & Objectives', level=1)
    doc.add_paragraph(
        'As Australia transitions towards Net Zero, the gas market faces a complex "dual challenge":'
    )
    p = doc.add_paragraph(style='List Bullet')
    p.add_run('Basin Depletion: ').bold = True
    p.add_run('The rapid decline of traditional supply sources, particularly the Gippsland Basin in the Bass Strait.')
    p.add_row = doc.add_paragraph(style='List Bullet')
    p.add_row.add_run('Demand Transition: ').bold = True
    p.add_row.add_run('Structural demand reduction via electrification of residential heating, offset by persistent industrial needs and export commitments.')
    
    doc.add_paragraph(
        'The objective of the model is to minimize total system cost, which is the sum of production costs, '
        'transportation tariffs, capital expenditure for new builds, and the economic cost of unserved demand.'
    )

    # Data Sources
    doc.add_heading('3. Key Data Sources & Assumptions', level=1)
    doc.add_paragraph(
        'The model is grounded in data from the Australian Energy Market Operator (AEMO) 2026 forecasting cycle.'
    )

    doc.add_heading('3.1 Demand Traces (Step Change Scenario)', level=2)
    doc.add_paragraph(
        'Annual demand profiles are based on the AEMO 2026 GSOO Step Change scenario. In 2026, Queensland LNG export '
        'demand is calibrated to ~3,650 TJ/day (~1,250 PJ/year). Long-term annual decline rates applied include:'
    )
    table = doc.add_table(rows=1, cols=3)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Node / Region'
    hdr_cells[1].text = 'Growth/Decline Rate'
    hdr_cells[2].text = 'Rationale'
    
    data = [
        ['Melbourne', '-4.0% p.a.', 'Victorian Gas Substitution Roadmap'],
        ['Sydney', '-3.0% p.a.', 'NSW residential electrification'],
        ['Brisbane', '-1.0% p.a.', 'Persistent industrial baseload'],
        ['LNG Export Cluster', '3,650 TJ/day', 'AEMO 2026 GSOO Benchmark']
    ]
    for region, rate, reason in data:
        row_cells = table.add_row().cells
        row_cells[0].text = region
        row_cells[1].text = rate
        row_cells[2].text = reason

    doc.add_paragraph('\nRelevant Links:')
    doc.add_paragraph('AEMO 2026 GSOO Report Data:', style='List Bullet')
    doc.add_paragraph('https://www.aemo.com.au/-/media/files/gas/national_planning_and_forecasting/gsoo/2026/2026-gas-statement-of-opportunities-report-figures-and-data.xlsx', style='Caption')
    doc.add_paragraph('AEMO 2026 GSOO Supply Data:', style='List Bullet')
    doc.add_paragraph('https://www.aemo.com.au/-/media/files/gas/national_planning_and_forecasting/gsoo/2026/2026-gas-statement-of-opportunities-supply-data.xlsx', style='Caption')

    doc.add_heading('3.2 Supply Basin Dynamics', level=2)
    doc.add_paragraph(
        'Production capacities reflect the 2026 GSOO benchmarks. Decline rates for southern basins have been '
        'accelerated to reflect recent depletion reports.'
    )
    s_table = doc.add_table(rows=1, cols=3)
    s_table.rows[0].cells[0].text = 'Basin'
    s_table.rows[0].cells[1].text = 'Capacity (TJ/d)'
    s_table.rows[0].cells[2].text = 'Annual Decline Rate'
    s_data = [
        ['Gippsland (Bass Strait)', '766', '12.0%'],
        ['Moomba (Cooper Basin)', '400', '3.0%'],
        ['Surat (QLD CSG)', '4,000', '1.0%']
    ]
    for basin, cap, decline in s_data:
        r = s_table.add_row().cells
        r[0].text = basin
        r[1].text = cap
        r[2].text = decline

    # Technical implementation
    doc.add_heading('4. Technical Implementation', level=1)
    
    doc.add_heading('4.1 Mathematical Optimization', level=2)
    doc.add_paragraph(
        'The model uses Mixed-Integer Linear Programming (MILP). The core constraints include nodal mass balance '
        '(Supply + Inflow = Demand + Outflow), pipeline capacities, and storage continuity.'
    )

    doc.add_heading('4.2 Storage Modeling', level=2)
    doc.add_paragraph(
        'Gas storage facilities (Iona, Moomba, Silver Springs) are modeled as "stateful" nodes. The inventory '
        'at the end of day (t) must equal the inventory from day (t-1) plus injections minus withdrawals. This '
        'allows the model to "pre-fill" storage in summer to meet the winter peaks triggered by the Southern Winter Stress scenarios.'
    )

    doc.add_heading('4.3 Multi-Year Path Dependency', level=2)
    doc.add_paragraph(
        'A unique feature of this simulation is the "already_built" logic. Infrastructure decisions are binary (0 or 1). '
        'Once the optimizer decides that a project like the Port Kembla LNG Import Terminal is economically necessary, '
        'that decision is locked in for all subsequent years in the 2050 sequence.'
    )

    # Pricing
    doc.add_heading('5. Price Discovery', level=1)
    doc.add_paragraph(
        'Nodal prices are derived from the marginal cost of supply (shadow prices). When a pipeline is uncongested, '
        'nodal prices are coupled (differing only by transport costs). When a pipeline hits its capacity limit, the '
        'nodal prices "decouple," reflecting local scarcity or surplus.'
    )

    # Save
    doc.save('gas_market_model/Gas_Market_Model_Documentation.docx')
    print("Created Gas_Market_Model_Documentation.docx")

if __name__ == "__main__":
    create_docs()
